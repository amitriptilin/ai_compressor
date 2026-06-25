import logging 

from docx import Document
from constants import CHUNK_SEPARATOR

logger = logging.getLogger(__name__)

def save_as_word(text: str, output_path: str) -> None:
    doc = Document()
    sections = text.split(CHUNK_SEPARATOR)
    
    for section in sections:
        lines = section.strip().split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("-"):
                doc.add_paragraph(line[1:].strip(), style = 'List Bullet')
            else:
                doc.add_paragraph(line)
            
        if len(sections) > 1 and section != sections[-1]:
            doc.add_page_break()
    
    doc.save(output_path)
    logger.info(f"Сохранено в DOCX: {output_path}")